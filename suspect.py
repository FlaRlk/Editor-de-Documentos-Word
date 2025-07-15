import os
import logging
from docx import Document
import glob
import customtkinter as ctk
import threading
import time
import sys
import win32com.client
import subprocess
import psutil
from datetime import datetime

# Config
APP_NAME = "Suspect Word Edit"
APP_VERSION = "2.0.0"
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue")

def setup_logging():
    
    logs_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
    os.makedirs(logs_folder, exist_ok=True)
    
    
    current_date = datetime.now()
    log_filename = os.path.join(logs_folder, f'suspect-{current_date.strftime("%d-%m-%Y")}.log')
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_filename, encoding='utf-8'),
            logging.StreamHandler(sys.stdout)
        ]
    )
    return log_filename

class TextReplacementFrame(ctk.CTkFrame):
    def __init__(self, master, index, on_delete=None, **kwargs):
        parent = master
        while parent and not hasattr(parent, 'colors'):
            parent = parent.master
        
        self.colors = parent.colors if parent and hasattr(parent, 'colors') else {
            "bg_dark": "#121212",
            "bg_medium": "#1E1E1E",
            "bg_light": "#2D2D2D",
            "accent": "#8A2BE2",
            "accent_hover": "#9932CC",
            "accent_light": "#BA55D3",
            "text_primary": "#FFFFFF",
            "text_secondary": "#B3B3B3",
            "success": "#4CAF50",
            "warning": "#FFA726",
            "error": "#EF5350",
            "border": "#333333"
        }
        
        super().__init__(
            master,
            fg_color=self.colors["bg_light"],
            border_color=self.colors["border"],
            border_width=1,
            **kwargs
        )
        
        self.index = index
        self.on_delete = on_delete
        
        self.header = ctk.CTkFrame(
            self,
            fg_color=self.colors["bg_medium"],
            height=35
        )
        self.header.pack(fill="x", padx=1, pady=(1, 5))
        self.header.pack_propagate(False)
        
        self.title_label = ctk.CTkLabel(
            self.header, 
            text=f"Substituição #{index}", 
            font=("Segoe UI", 12, "bold"),
            text_color=self.colors["accent_light"]
        )
        self.title_label.pack(side="left", padx=10, pady=5)
        
        if on_delete and index > 1:
            self.delete_btn = ctk.CTkButton(
                self.header,
                text="✕",
                width=25,
                height=25,
                fg_color="transparent",
                hover_color=self.colors["error"],
                text_color=self.colors["text_secondary"],
                command=self._on_delete,
                font=("Segoe UI", 12)
            )
            self.delete_btn.pack(side="right", padx=5, pady=5)
        
        self.fields_container = ctk.CTkFrame(
            self,
            fg_color="transparent"
        )
        self.fields_container.pack(fill="x", padx=10, pady=(0, 10))
        
        self.find_label = ctk.CTkLabel(
            self.fields_container,
            text="Texto Original",
            font=("Segoe UI", 11),
            text_color=self.colors["text_secondary"]
        )
        self.find_label.pack(anchor="w", pady=(5, 0))
        
        self.find_text = ctk.CTkEntry(
            self.fields_container,
            placeholder_text="Digite o texto que deseja encontrar",
            height=35,
            font=("Segoe UI", 12),
            fg_color=self.colors["bg_medium"],
            border_color=self.colors["border"],
            border_width=1,
            placeholder_text_color=self.colors["text_secondary"]
        )
        self.find_text.pack(fill="x", pady=(0, 10))
        
        self.replace_label = ctk.CTkLabel(
            self.fields_container,
            text="Novo Texto",
            font=("Segoe UI", 11),
            text_color=self.colors["text_secondary"]
        )
        self.replace_label.pack(anchor="w", pady=(5, 0))
        
        self.replace_text = ctk.CTkEntry(
            self.fields_container,
            placeholder_text="Digite o texto que substituirá o original",
            height=35,
            font=("Segoe UI", 12),
            fg_color=self.colors["bg_medium"],
            border_color=self.colors["border"],
            border_width=1,
            placeholder_text_color=self.colors["text_secondary"]
        )
        self.replace_text.pack(fill="x", pady=(0, 5))
        
        self.case_sensitive_var = ctk.BooleanVar(value=False)
        self.case_sensitive_checkbox = ctk.CTkCheckBox(
            self.fields_container,
            text="Diferenciar maiúsculas/minúsculas",
            variable=self.case_sensitive_var,
            font=("Segoe UI", 11),
            text_color=self.colors["text_secondary"],
            fg_color=self.colors["accent"],
            hover_color=self.colors["accent_hover"],
            checkbox_width=20,
            checkbox_height=20
        )
        self.case_sensitive_checkbox.pack(anchor="w", pady=(5, 0))
    
    def _on_delete(self):
        if self.on_delete:
            self.on_delete(self)
    
    def get_values(self):
        return {
            "find": self.find_text.get().strip(),
            "replace": self.replace_text.get().strip(),
            "case_sensitive": self.case_sensitive_var.get()
        }

class ModernWordProcessor(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME)
        self.geometry("1200x800")
        self.grid_columnconfigure(0, weight=3)
        self.grid_columnconfigure(1, weight=2)
        self.grid_rowconfigure(0, weight=1)
        
        self.processing = False
        self.paused = False
        self.stop_requested = False
        self.current_thread = None
        
        self.log_filename = setup_logging()
        
        self.colors = {
            "bg_dark": "#121212",
            "bg_medium": "#1E1E1E",
            "bg_light": "#2D2D2D",
            "accent": "#8A2BE2",
            "accent_hover": "#9932CC",
            "accent_light": "#BA55D3",
            "text_primary": "#FFFFFF",
            "text_secondary": "#B3B3B3",
            "success": "#4CAF50",
            "warning": "#FFA726",
            "error": "#EF5350",
            "border": "#333333"
        }
        
        self.configure(fg_color=self.colors["bg_dark"])
        
        
        try:
            icon_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "icon.ico")
            if os.path.exists(icon_path):
                self.iconbitmap(icon_path)
            else:
                logging.warning("Arquivo de ícone não encontrado: %s", icon_path)
        except Exception as e:
            logging.error("Erro ao carregar o ícone: %s", str(e))
        
        self.left_frame = ctk.CTkFrame(
            self,
            fg_color=self.colors["bg_medium"],
            border_color=self.colors["border"],
            border_width=1
        )
        self.left_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        
        self.right_frame = ctk.CTkFrame(
            self,
            fg_color=self.colors["bg_medium"],
            border_color=self.colors["border"],
            border_width=1
        )
        self.right_frame.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")
        
        self.word_folder = None
        self.pdf_folder = None
        self.word_app = None
        self.replacement_frames = []
        
        self.search_across_paragraphs = ctk.BooleanVar(value=False)
        
        self.setup_left_frame()
        self.setup_right_frame()
        
        self.protocol("WM_DELETE_WINDOW", self.on_closing)

    def setup_left_frame(self):
        self.left_frame.grid_columnconfigure(0, weight=1)
        
        title_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color="transparent"
        )
        title_frame.grid(row=0, column=0, pady=(20, 5), padx=20, sticky="ew")
        
        title_label = ctk.CTkLabel(
            title_frame,
            text="Suspect Word Edit",
            font=("Segoe UI", 28, "bold"),
            text_color=self.colors["accent"]
        )
        title_label.pack(side="left")
        
        version_label = ctk.CTkLabel(
            title_frame,
            text=f"v{APP_VERSION}",
            font=("Segoe UI", 12),
            text_color=self.colors["text_secondary"]
        )
        version_label.pack(side="left", padx=(5, 0), pady=(8, 0))
        
        instructions_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color=self.colors["bg_light"],
            border_color=self.colors["border"],
            border_width=1
        )
        instructions_frame.grid(row=1, column=0, padx=20, pady=10, sticky="ew")
        
        instructions_text = "Este programa permite substituir textos em documentos Word e convertê-los para PDF."
        instructions_label = ctk.CTkLabel(
            instructions_frame,
            text=instructions_text,
            font=("Segoe UI", 12),
            text_color=self.colors["text_secondary"],
            wraplength=500
        )
        instructions_label.pack(padx=15, pady=15)
        
        self.config_label = ctk.CTkLabel(
            self.left_frame,
            text="Configurações de Substituição",
            font=("Segoe UI", 16, "bold"),
            text_color=self.colors["accent_light"]
        )
        self.config_label.grid(row=2, column=0, pady=(20, 5), sticky="w", padx=20)
        
        advanced_options_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color=self.colors["bg_light"],
            border_color=self.colors["border"],
            border_width=1
        )
        advanced_options_frame.grid(row=3, column=0, padx=20, pady=5, sticky="ew")
        
        advanced_options_title = ctk.CTkLabel(
            advanced_options_frame,
            text="Opções Avançadas",
            font=("Segoe UI", 14, "bold"),
            text_color=self.colors["accent_light"]
        )
        advanced_options_title.pack(anchor="w", padx=15, pady=(10, 5))
        
        self.cross_paragraph_checkbox = ctk.CTkCheckBox(
            advanced_options_frame,
            text="Buscar textos divididos entre parágrafos (mais lento)",
            variable=self.search_across_paragraphs,
            font=("Segoe UI", 12),
            text_color=self.colors["text_secondary"],
            fg_color=self.colors["accent"],
            hover_color=self.colors["accent_hover"],
            checkbox_width=20,
            checkbox_height=20
        )
        self.cross_paragraph_checkbox.pack(anchor="w", padx=15, pady=(5, 15))
        
        replacements_outer_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color=self.colors["bg_light"],
            border_color=self.colors["border"],
            border_width=1
        )
        replacements_outer_frame.grid(row=4, column=0, padx=20, pady=5, sticky="nsew")
        replacements_outer_frame.grid_columnconfigure(0, weight=1)
        replacements_outer_frame.grid_rowconfigure(0, weight=1)
        
        self.replacements_container = ctk.CTkScrollableFrame(
            replacements_outer_frame,
            fg_color="transparent"
        )
        self.replacements_container.grid(row=0, column=0, padx=1, pady=1, sticky="nsew")
        self.left_frame.grid_rowconfigure(4, weight=1)
        
        self.add_replacement_frame()
        
        self.add_btn_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color="transparent"
        )
        self.add_btn_frame.grid(row=5, column=0, padx=20, pady=10, sticky="ew")
        
        self.add_btn = ctk.CTkButton(
            self.add_btn_frame,
            text="+ Adicionar Substituição",
            command=self.add_replacement_frame,
            fg_color=self.colors["bg_light"],
            hover_color=self.colors["accent"],
            height=35,
            border_color=self.colors["border"],
            border_width=1
        )
        self.add_btn.pack(side="left", padx=5)
        
        self.control_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color=self.colors["bg_light"],
            border_color=self.colors["border"],
            border_width=1
        )
        self.control_frame.grid(row=6, column=0, padx=20, pady=10, sticky="ew")
        
        self.process_button = ctk.CTkButton(
            self.control_frame,
            text="▶ Iniciar Processamento",
            command=self.start_processing,
            font=("Segoe UI", 14, "bold"),
            height=40,
            fg_color=self.colors["accent"],
            hover_color=self.colors["accent_hover"]
        )
        self.process_button.pack(side="left", padx=5, pady=5, expand=True, fill="x")
        
        self.pause_button = ctk.CTkButton(
            self.control_frame,
            text="⏸",
            command=self.pause_processing,
            font=("Segoe UI", 14, "bold"),
            width=50,
            height=40,
            fg_color=self.colors["warning"],
            hover_color="#E69500",
            state="disabled"
        )
        self.pause_button.pack(side="left", padx=2, pady=5)
        
        self.stop_button = ctk.CTkButton(
            self.control_frame,
            text="⏹",
            command=self.stop_processing,
            font=("Segoe UI", 14, "bold"),
            width=50,
            height=40,
            fg_color=self.colors["error"],
            hover_color="#D32F2F",
            state="disabled"
        )
        self.stop_button.pack(side="left", padx=2, pady=5)
        
        progress_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color=self.colors["bg_light"],
            border_color=self.colors["border"],
            border_width=1
        )
        progress_frame.grid(row=7, column=0, padx=20, pady=10, sticky="ew")
        
        self.progress_label = ctk.CTkLabel(
            progress_frame,
            text="Pronto para começar!",
            font=("Segoe UI", 12),
            text_color=self.colors["text_secondary"]
        )
        self.progress_label.pack(pady=(10, 0), padx=10)
        
        self.progress_bar = ctk.CTkProgressBar(
            progress_frame,
            progress_color=self.colors["accent"],
            height=15
        )
        self.progress_bar.pack(pady=10, padx=10, fill="x")
        self.progress_bar.set(0)
        
        log_info_frame = ctk.CTkFrame(
            self.left_frame,
            fg_color="transparent"
        )
        log_info_frame.grid(row=8, column=0, padx=20, pady=5, sticky="ew")
        
        
        log_path = self.log_filename
        if len(log_path) > 50:
            log_path = "..." + log_path[-47:]
            
        self.log_info_label = ctk.CTkLabel(
            log_info_frame,
            text=f"Log: {log_path}",
            font=("Segoe UI", 10),
            text_color=self.colors["text_secondary"]
        )
        self.log_info_label.pack(anchor="w")
        
    def setup_right_frame(self):
        log_title_frame = ctk.CTkFrame(
            self.right_frame,
            fg_color="transparent"
        )
        log_title_frame.pack(pady=(20, 10), fill="x")
        
        log_title = ctk.CTkLabel(
            log_title_frame,
            text="Logs",
            font=("Segoe UI", 16, "bold"),
            text_color=self.colors["accent"]
        )
        log_title.pack(side="left", padx=20)
        
        log_container = ctk.CTkFrame(
            self.right_frame,
            fg_color=self.colors["bg_light"],
            border_color=self.colors["border"],
            border_width=1
        )
        log_container.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        
        self.log_text = ctk.CTkTextbox(
            log_container,
            font=("Consolas", 12),
            wrap="word",
            fg_color=self.colors["bg_light"],
            text_color=self.colors["text_secondary"],
            border_color=self.colors["border"],
            border_width=1
        )
        self.log_text.pack(fill="both", expand=True, padx=5, pady=5)
    
    def add_replacement_frame(self):
        index = len(self.replacement_frames) + 1
        frame = TextReplacementFrame(
            self.replacements_container,
            index,
            on_delete=self.delete_replacement_frame
        )
        frame.pack(fill="x", padx=5, pady=5, expand=False)
        self.replacement_frames.append(frame)
    
    def delete_replacement_frame(self, frame):
        if len(self.replacement_frames) > 1:
            self.replacement_frames.remove(frame)
            frame.destroy()
            
            for i, f in enumerate(self.replacement_frames, 1):
                f.index = i
                f.title_label.configure(text=f"Substituição #{i}")
    
    def start_processing(self):
        if not self.processing:
            self.processing = True
            self.paused = False
            self.stop_requested = False
            self.process_button.configure(text="▶ Processando...", state="disabled")
            self.pause_button.configure(state="normal")
            self.stop_button.configure(state="normal")
            self.log_text.delete("1.0", "end")
            self.log("Iniciando o processamento dos arquivos...", "info")
            self.current_thread = threading.Thread(target=self.process_files, daemon=True)
            self.current_thread.start()
    
    def pause_processing(self):
        if self.processing:
            if not self.paused:
                self.paused = True
                self.pause_button.configure(text="▶")
                self.process_button.configure(text="⏸ Pausado")
                self.log("Processamento pausado", "warning")
            else:
                self.paused = False
                self.pause_button.configure(text="⏸")
                self.process_button.configure(text="▶ Processando...")
                self.log("Processamento retomado", "info")
    
    def stop_processing(self):
        if self.processing:
            self.stop_requested = True
            self.paused = False
            self.log("Solicitação de parada recebida...", "warning")
            self.process_button.configure(text="⏹ Parando...")
    
    def reset_buttons(self):
        self.processing = False
        self.paused = False
        self.stop_requested = False
        self.process_button.configure(text="▶ Iniciar Processamento", state="normal")
        self.pause_button.configure(text="⏸", state="disabled")
        self.stop_button.configure(state="disabled")
    
    def log(self, message, level="info"):
        colors = {
            "info": self.colors["text_primary"],
            "success": self.colors["success"],
            "error": self.colors["error"],
            "warning": self.colors["warning"]
        }
        
        timestamp = time.strftime("%H:%M:%S")
        
        if level == "success":
            prefix = "✓"
            emoji = "✅ "
        elif level == "error":
            prefix = "×"
            emoji = "❌ "
        elif level == "warning":
            prefix = "!"
            emoji = "⚠️ "
        else:
            prefix = "→"
            emoji = "📝 "
            
        formatted_message = f"[{timestamp}] {emoji}{message}\n"
        self.log_text.insert("end", formatted_message)
        self.log_text.see("end")
        
        end_index = self.log_text.index("end-1c")
        start_index = f"{float(end_index) - 1:.1f}"
        self.log_text.tag_add(level, start_index, end_index)
        self.log_text.tag_config(level, foreground=colors[level])
        
        logging.info(f"{prefix} {message}")
    
    def update_status(self, message, progress=None):
        self.progress_label.configure(text=message)
        if progress is not None:
            self.progress_bar.set(progress)
    
    def wait_if_paused(self):
        if self.paused:
            self.log("Processamento pausado", "warning")
            while self.paused and not self.stop_requested:
                time.sleep(0.1)
                self.update()
            if not self.stop_requested:
                self.log("Processamento continuado", "info")
    
    def normalize_text(self, text):
        return ' '.join(text.split()).lower()
    
    def kill_word_processes(self):
        try:
            for proc in psutil.process_iter(['pid', 'name']):
                try:
                    if proc.info['name'] and 'winword.exe' in proc.info['name'].lower():
                        proc.kill()
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    pass
            time.sleep(0.5)
        except Exception as e:
            self.log(f"Erro ao finalizar processos Word: {str(e)}", "warning")
    
    def convert_to_pdf(self, word_path):
        if self.stop_requested:
            return False
            
        self.wait_if_paused()
        
        self.kill_word_processes()
        time.sleep(1)
        
        pdf_path = os.path.join(
            self.pdf_folder,
            os.path.splitext(os.path.basename(word_path))[0] + '.pdf'
        )
        
        if os.path.exists(pdf_path):
            try:
                os.remove(pdf_path)
                self.log(f"PDF antigo removido: {os.path.basename(pdf_path)}", "info")
            except Exception as e:
                self.log(f"Não foi possível remover o PDF antigo: {str(e)}", "warning")
                return False

        try:
            self.log(f"Convertendo para PDF usando PowerShell: {os.path.basename(word_path)}", "info")
            
            ps_script = f"""
            $wordApp = New-Object -ComObject Word.Application
            $wordApp.Visible = $false
            $wordApp.DisplayAlerts = 0
            $doc = $wordApp.Documents.Open("{os.path.abspath(word_path).replace('\\', '\\\\')}")
            $pdf_path = "{os.path.abspath(pdf_path).replace('\\', '\\\\')}"
            $doc.SaveAs([ref] $pdf_path, [ref] 17)
            $doc.Close()
            $wordApp.Quit()
            [System.Runtime.Interopservices.Marshal]::ReleaseComObject($doc) | Out-Null
            [System.Runtime.Interopservices.Marshal]::ReleaseComObject($wordApp) | Out-Null
            [System.GC]::Collect()
            [System.GC]::WaitForPendingFinalizers()
            """
            
            ps_path = os.path.join(os.path.dirname(word_path), "convert_to_pdf.ps1")
            with open(ps_path, "w") as ps_file:
                ps_file.write(ps_script)
            
            import subprocess
            result = subprocess.run(
                ["powershell", "-ExecutionPolicy", "Bypass", "-File", ps_path],
                capture_output=True,
                text=True
            )
            
            try:
                os.remove(ps_path)
            except:
                pass
            
            if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                self.log(f"PDF gerado com sucesso: {os.path.basename(pdf_path)}", "success")
                return True
            else:
                self.log(f"Não foi possível gerar o PDF", "error")
                return False
                
        except Exception as e:
            self.log(f"Erro ao gerar PDF: {str(e)}", "error")
            self.kill_word_processes()
            return False
    
    def process_document(self, file_path):
        if self.stop_requested:
            return False
            
        self.wait_if_paused()
        
        try:
            self.log(f"Analisando documento: {os.path.basename(file_path)}", "info")
            doc = Document(file_path)
            modified = False
            
            replacements = []
            for frame in self.replacement_frames:
                values = frame.get_values()
                if values["find"] and values["replace"]:
                    replacements.append(values)
            
            if not replacements:
                self.log("Nenhuma substituição configurada", "warning")
                return False
            
            search_across_paragraphs = self.search_across_paragraphs.get()
            if search_across_paragraphs:
                self.log("Modo de busca entre parágrafos ativado", "info")
                if self.process_document_across_paragraphs(doc, replacements):
                    modified = True
            
            for section in doc.sections:
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        if self.process_paragraph_text(paragraph, replacements, "cabeçalho"):
                            modified = True
                
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if self.process_paragraph_text(paragraph, replacements, "cabeçalho"):
                                    modified = True
            
            footer = section.footer
            if footer:
                for paragraph in footer.paragraphs:
                    if self.process_paragraph_text(paragraph, replacements, "rodapé"):
                        modified = True
                
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if self.process_paragraph_text(paragraph, replacements, "rodapé"):
                                    modified = True
        
            for paragraph in doc.paragraphs:
                if self.process_paragraph_text(paragraph, replacements, "documento principal"):
                    modified = True
        
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if self.process_paragraph_text(paragraph, replacements, "documento principal"):
                                modified = True

            try:
                for shape in doc.inline_shapes:
                    if hasattr(shape, 'text_frame') and shape.text_frame:
                        for paragraph in shape.text_frame.paragraphs:
                            if self.process_paragraph_text(paragraph, replacements, "caixa de texto"):
                                modified = True
            except Exception as e:
                self.log(f"Aviso: Não foi possível processar algumas caixas de texto: {str(e)}", "warning")

            if modified:
                doc.save(file_path)
                self.log(f"Documento salvo com alterações: {os.path.basename(file_path)}", "success")
            else:
                self.log(f"Documento já está com as informações atualizadas: {os.path.basename(file_path)}", "info")
            
            return modified

        except Exception as e:
            self.log(f"Erro ao processar o documento: {str(e)}", "error")
            return False
            
    def process_document_across_paragraphs(self, doc, replacements):
        try:
            main_text = ""
            paragraph_positions = []
            start_pos = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if text:
                    main_text += text + "\n"
                    paragraph_positions.append((start_pos, start_pos + len(text), paragraph))
                    start_pos += len(text) + 1
            
            modified = False
            for replacement in replacements:
                find_text = replacement["find"]
                replace_text = replacement["replace"]
                case_sensitive = replacement.get("case_sensitive", False)
                
                if not find_text:
                    continue
                
                remaining_text = main_text
                offset = 0
                
                while True:
                    if case_sensitive:
                        pos = remaining_text.find(find_text)
                    else:
                        pos = remaining_text.lower().find(find_text.lower())
                        
                    if pos == -1:
                        break
                    
                    abs_pos = offset + pos
                    abs_end = abs_pos + len(find_text)
                    
                    affected_paragraphs = []
                    for start, end, paragraph in paragraph_positions:
                        if (start <= abs_pos < end) or (start < abs_end <= end) or (abs_pos <= start and abs_end >= end):
                            affected_paragraphs.append((start, end, paragraph))
                    
                    if len(affected_paragraphs) > 1:
                        self.log(f"Texto dividido entre {len(affected_paragraphs)} parágrafos: '{find_text}' → '{replace_text}'", "info")
                        
                        for start, end, paragraph in affected_paragraphs:
                            para_start = max(abs_pos - start, 0)
                            para_end = min(abs_end - start, end - start)
                            
                            original_text = paragraph.text
                            
                            if para_start == 0 and para_end == len(original_text):
                                if affected_paragraphs.index((start, end, paragraph)) == 0:
                                    new_text = replace_text
                                else:
                                    new_text = ""
                            else:
                                if affected_paragraphs.index((start, end, paragraph)) == 0:
                                    new_text = original_text[:para_start] + replace_text
                                elif affected_paragraphs.index((start, end, paragraph)) == len(affected_paragraphs) - 1:
                                    new_text = original_text[para_end:]
                                else:
                                    new_text = ""
                            
                            for i in range(len(paragraph.runs)-1, -1, -1):
                                p = paragraph._p
                                p.remove(paragraph.runs[i]._r)
                            
                            new_run = paragraph.add_run(new_text)
                            if paragraph.runs:
                                self.copy_run_formatting(paragraph.runs[0], new_run)
                            modified = True
                    
                    offset += pos + len(find_text)
                    remaining_text = remaining_text[pos + len(find_text):]
                    
                    main_text = ""
                    paragraph_positions = []
                    start_pos = 0
                    
                    for paragraph in doc.paragraphs:
                        text = paragraph.text
                        if text:
                            main_text += text + "\n"
                            paragraph_positions.append((start_pos, start_pos + len(text), paragraph))
                            start_pos += len(text) + 1
                    
                    remaining_text = main_text[offset:]
            
            return modified
        except Exception as e:
            self.log(f"Erro ao processar texto entre parágrafos: {str(e)}", "error")
            return False

    def text_contains(self, text, search_text, case_sensitive=False):
        if not case_sensitive:
            return search_text.lower() in text.lower()
        return search_text in text
            
    def replace_text(self, text, search_text, replace_text, case_sensitive=False):
        if not case_sensitive:
            result = ""
            remaining = text
            lower_search = search_text.lower()
            
            while remaining:
                lower_remaining = remaining.lower()
                pos = lower_remaining.find(lower_search)
                
                if pos == -1:
                    result += remaining
                    break
                    
                result += remaining[:pos]
                result += replace_text
                remaining = remaining[pos + len(search_text):]
                
            return result
        else:
            return text.replace(search_text, replace_text)
    
    def process_paragraph_text(self, paragraph, replacements, location="texto"):
        modified = False
        full_text = paragraph.text
        
        if not full_text:
            return False
            
        original_runs = list(paragraph.runs)
        current_pos = 0
        
        for replacement in replacements:
            find_text = replacement["find"]
            replace_text = replacement["replace"]
            case_sensitive = replacement.get("case_sensitive", False)
            
            if not find_text:
                continue
                
            if self.text_contains(full_text, find_text, case_sensitive):
                text_changed = False
                new_runs = []
                
                for run in original_runs:
                    run_text = run.text
                    if not run_text:
                        new_runs.append(run)
                        continue
                    
                    while run_text:
                        if case_sensitive:
                            pos = run_text.find(find_text)
                        else:
                            pos = run_text.lower().find(find_text.lower())
                            
                        if pos == -1:
                            new_run = paragraph.add_run(run_text)
                            self.copy_run_formatting(run, new_run)
                            new_runs.append(new_run)
                            break
                        
                        if pos > 0:
                            before_run = paragraph.add_run(run_text[:pos])
                            self.copy_run_formatting(run, before_run)
                            new_runs.append(before_run)
                        
                        replace_run = paragraph.add_run(replace_text)
                        self.copy_run_formatting(run, replace_run)
                        new_runs.append(replace_run)
                        
                        run_text = run_text[pos + len(find_text):]
                        text_changed = True
                        modified = True
                
                if text_changed:
                    for run in original_runs:
                        paragraph._p.remove(run._r)
                    original_runs = new_runs
                    
                    count = full_text.lower().count(find_text.lower()) if not case_sensitive else full_text.count(find_text)
                    self.log(f"Texto substituído no {location}: '{find_text}' → '{replace_text}' ({count} ocorrência{'s' if count > 1 else ''})", "success")
        
        return modified

    def copy_run_formatting(self, source_run, target_run):
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.color.rgb = source_run.font.color.rgb if source_run.font.color.rgb else None
        target_run.style = source_run.style

    def process_files(self):
        try:
            has_valid_replacement = False
            for frame in self.replacement_frames:
                values = frame.get_values()
                if values["find"] and values["replace"]:
                    has_valid_replacement = True
                    break
            
            if not has_valid_replacement:
                self.log("Por favor, configure pelo menos uma substituição de texto", "warning")
                self.update_status("Configuração incompleta")
                return
            
            self.word_folder = ctk.filedialog.askdirectory(
                title="1. Selecione a pasta com seus arquivos Word"
            )
            if not self.word_folder:
                self.update_status("Operação cancelada")
                return
                
            self.pdf_folder = ctk.filedialog.askdirectory(
                title="2. Selecione onde deseja salvar os PDFs"
            )
            if not self.pdf_folder:
                self.update_status("Operação cancelada")
                return
                
            if os.path.normpath(self.word_folder) == os.path.normpath(self.pdf_folder):
                self.log("Por favor, escolha pastas diferentes para os arquivos Word e PDF", "warning")
                self.update_status("Mesma pasta selecionada")
                return
                
            word_files = glob.glob(os.path.join(self.word_folder, "*.docx"))
            if not word_files:
                self.update_status("Nenhum arquivo Word encontrado")
                self.log("Não encontrei nenhum arquivo .docx na pasta selecionada", "warning")
                return
            
            word_files = [f for f in word_files if not os.path.basename(f).startswith("~$")]
                
            total_files = len(word_files)
            successful_word = 0
            successful_pdf = 0
            already_updated = 0
            no_changes_needed = 0
            
            self.process_button.configure(state="disabled")
            self.pause_button.configure(state="normal")
            self.stop_button.configure(state="normal")
            
            self.log(f"Iniciando processamento de {total_files} documentos", "info")
            
            self.kill_word_processes()
            time.sleep(1)
                
            for i, file in enumerate(word_files, 1):
                if not self.processing:
                    break
                    
                progress = i / total_files
                self.update_status(f"Processando arquivo {i} de {total_files}: {os.path.basename(file)}", progress)
                
                was_modified = self.process_document(file)
                
                if was_modified:
                    successful_word += 1
                    if self.convert_to_pdf(file):
                        successful_pdf += 1
                    else:
                        self.log(f"Falha na geração do PDF: {os.path.basename(file)}", "error")
                else:
                    log_text = self.log_text.get("1.0", "end")
                    if "já está com as informações atualizadas" in log_text:
                        already_updated += 1
                        if self.convert_to_pdf(file):
                            successful_pdf += 1
                        else:
                            self.log(f"Falha na geração do PDF: {os.path.basename(file)}", "error")
                    else:
                        no_changes_needed += 1
                        
                time.sleep(0.5)
                
            self.kill_word_processes()
            self.log(f"Processamento Word finalizado", "info")
            
            self.reset_buttons()
            
            final_message = (
                f"Processamento concluído!\n"
                f"• Documentos atualizados: {successful_word}/{total_files}\n"
                f"• PDFs gerados: {successful_pdf}/{total_files}\n"
                f"• Já estavam atualizados: {already_updated}/{total_files}\n"
                f"• Sem necessidade de mudanças: {no_changes_needed}/{total_files}"
            )
            
            self.update_status(final_message, 1.0)
            self.log(final_message, "success")
            
        except Exception as e:
            self.log(f"Erro no processamento: {str(e)}", "error")
            self.update_status("Erro no processamento")
            self.kill_word_processes()
            self.reset_buttons()
    
    def on_closing(self):
        if self.processing:
            self.stop_requested = True
            self.log("Finalizando aplicação...", "info")
            time.sleep(1)
        
        if self.word_app:
            try:
                self.word_app.Quit()
            except:
                pass
        
        self.kill_word_processes()
        self.destroy()

if __name__ == "__main__":
    
    if getattr(sys, 'frozen', False):
       
        application_path = os.path.dirname(sys.executable)
    else:
        application_path = os.path.dirname(os.path.abspath(__file__))
    
    os.chdir(application_path)
    
    app = ModernWordProcessor()
    app.mainloop()